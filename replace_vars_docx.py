import json
import sys
import docx # Import the module, not the function/class
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table

# The actual Document class is located here:
from docx.document import Document as OriginalDocumentClass # Import the class type for isinstance checks

def iter_block_items(parent):
    """
    Itère sur tous les paragraphes et tableaux dans un document,
    une cellule ou un en-tête/pied de page.
    """
    # Use the specific Document class imported above for reliable type checking
    if isinstance(parent, OriginalDocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif hasattr(parent, 'element'):  # Header/Footer
        parent_elm = parent.element
    else:
        return
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def remplacer_texte_complet(paragraph, variables, delimiteur_gauche, delimiteur_droit):
    """
    Remplace les variables même si elles sont éclatées sur plusieurs runs.
    """
    texte_complet = "".join(run.text for run in paragraph.runs)
    nouveau_texte = texte_complet
    for cle, valeur in variables.items():
        placeholder = f"{delimiteur_gauche}{cle}{delimiteur_droit}"
        nouveau_texte = nouveau_texte.replace(placeholder, str(valeur))

    if nouveau_texte != texte_complet:
        # Clear all existing runs
        for run in paragraph.runs:
            run.text = ""
        # Write the new text into the first run
        if paragraph.runs:
            paragraph.runs[0].text = nouveau_texte
        else:
            # Handle case of empty paragraph
            paragraph.add_run(nouveau_texte)


def remplacer_variables_docx(doc_path, json_path, output_path,
                             delimiteur_gauche="{{", delimiteur_droit="}}"):
    # Charger les variables JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        variables = json.load(f)

    # Charger le document. docx.Document() is a function that returns a 
    # docx.document.Document object.
    doc = docx.Document(doc_path) 

    # Corps du document
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            remplacer_texte_complet(block, variables, delimiteur_gauche, delimiteur_droit)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for inner_block in iter_block_items(cell):
                        if isinstance(inner_block, Paragraph):
                            remplacer_texte_complet(inner_block, variables, delimiteur_gauche, delimiteur_droit)

    # En-têtes et pieds de page
    for section in doc.sections:
        # Check if header/footer exists before iterating
        if section.header:
            for block in iter_block_items(section.header):
                if isinstance(block, Paragraph):
                    remplacer_texte_complet(block, variables, delimiteur_gauche, delimiteur_droit)
        if section.footer:
            for block in iter_block_items(section.footer):
                if isinstance(block, Paragraph):
                    remplacer_texte_complet(block, variables, delimiteur_gauche, delimiteur_droit)

    # Sauvegarde du document
    doc.save(output_path)
    print(f"✅ Document généré avec succès : {output_path}")

# Ligne de commande
if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage : python replace_vars_docx.py <input.docx> <variables.json> <output.docx> [delimiteur_gauche] [delimiteur_droit]")
        sys.exit(1)

    docx_file = sys.argv[1]
    json_file = sys.argv[2]
    output_file = sys.argv[3]
    delimiteur_gauche = sys.argv[4] if len(sys.argv) > 4 else "{{"
    delimiteur_droit = sys.argv[5] if len(sys.argv) > 5 else "}}"

    remplacer_variables_docx(docx_file, json_file, output_file, delimiteur_gauche, delimiteur_droit)